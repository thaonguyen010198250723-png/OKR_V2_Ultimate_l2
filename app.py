import streamlit as st
import pandas as pd
from supabase import create_client, Client
import extra_streamlit_components as stx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import time
import uuid
import datetime

# =============================================================================
# 1. CẤU HÌNH & KẾT NỐI SUPABASE
# =============================================================================

st.set_page_config(
    page_title="Hệ thống Quản lý OKR (Supabase)",
    page_icon="🏫",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Ẩn Branding của Streamlit
st.markdown("""
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display:none;}
    </style>
    """, unsafe_allow_html=True)

# Khởi tạo Cookie Manager
cookie_manager = stx.CookieManager()

# Cấu hình Supabase (Credentials)
SUPABASE_URL = "https://iwobcnevhvqavonbjnnw.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Iml3b2JjbmV2aHZxYXZvbmJqbm53Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3Njk0NDI3MDMsImV4cCI6MjA4NTAxODcwM30.InEuVLSU3NBtbQg7yB0E9AI21LK73RWc8TcvPPvOvjw"

@st.cache_resource
def init_supabase():
    return create_client(SUPABASE_URL, SUPABASE_KEY)

supabase: Client = init_supabase()

# Định nghĩa Schema chuẩn (Để map dữ liệu khi DF rỗng)
SCHEMA = {
    'Users': ['Email', 'Password', 'Role', 'HoTen', 'Lop', 'EmailPH', 'SiSo'],
    'Periods': ['TenDot', 'TrangThai'],
    'OKRs': ['ID', 'Email', 'Lop', 'Dot', 'MucTieu', 'KetQuaThenChot', 
             'MucTieuSo', 'ThucDat', 'DonVi', 'TienDo', 'TrangThai', 
             'YeuCauXoa', 'NhanXet_GV', 'DiemHaiLong_PH', 'NhanXet_PH'],
    'FinalReviews': ['Email', 'Dot', 'NhanXet_CuoiKy', 'PhanHoi_PH', 'TrangThai_CuoiKy']
}

# Session State Init
if 'user' not in st.session_state:
    st.session_state.user = None
if 'active_expander' not in st.session_state:
    st.session_state.active_expander = None

# =============================================================================
# 2. XỬ LÝ DỮ LIỆU (DATA LAYER) - FIX ARROW ERROR
# =============================================================================

@st.cache_data(ttl=5) # Cache ngắn để data realtime hơn, giảm request thừa
def load_data(table_name):
    """
    Tải dữ liệu từ Supabase và ÉP KIỂU CHI TIẾT để tránh lỗi PyArrow.
    """
    try:
        response = supabase.table(table_name).select("*").execute()
        data = response.data
        
        # 1. Xử lý trường hợp data rỗng -> Trả về DataFrame đúng Schema
        if not data:
            return pd.DataFrame(columns=SCHEMA.get(table_name, []))
        
        df = pd.DataFrame(data)
        
        # 2. FIX LỖI PYARROW: Ép kiểu dữ liệu nghiêm ngặt
        
        # Nhóm cột String/UUID (Chuyển hết về string để tránh lỗi UUID object)
        text_cols = ['ID', 'Email', 'Password', 'Role', 'HoTen', 'Lop', 'EmailPH', 
                     'TenDot', 'TrangThai', 'MucTieu', 'KetQuaThenChot', 'DonVi', 
                     'NhanXet_GV', 'NhanXet_PH', 'NhanXet_CuoiKy', 'PhanHoi_PH', 
                     'TrangThai_CuoiKy']
        
        for col in text_cols:
            if col in df.columns:
                df[col] = df[col].astype(str).replace(['None', 'nan'], '')

        # Nhóm cột Số (Float) - Fill NaN bằng 0.0
        num_cols = ['SiSo', 'MucTieuSo', 'ThucDat', 'TienDo', 'DiemHaiLong_PH']
        for col in num_cols:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0.0)

        return df

    except Exception as e:
        st.error(f"Lỗi tải dữ liệu bảng {table_name}: {str(e)}")
        return pd.DataFrame(columns=SCHEMA.get(table_name, []))

def clear_cache():
    st.cache_data.clear()

def upsert_data(table_name, data_dict):
    """Thêm mới hoặc Cập nhật (Dựa trên Primary Key của Table)"""
    try:
        supabase.table(table_name).upsert(data_dict).execute()
        clear_cache()
        return True
    except Exception as e:
        st.error(f"Lỗi lưu dữ liệu: {e}")
        return False

def delete_data(table_name, column_name, value):
    """Xóa dòng dữ liệu"""
    try:
        supabase.table(table_name).delete().eq(column_name, value).execute()
        clear_cache()
        return True
    except Exception as e:
        st.error(f"Lỗi xóa dữ liệu: {e}")
        return False

# =============================================================================
# 3. UTILITIES
# =============================================================================

def calculate_progress(actual, target):
    try:
        t = float(target)
        a = float(actual)
        if t == 0: return 100.0 if a > 0 else 0.0
        return min((a / t) * 100.0, 100.0)
    except:
        return 0.0

def generate_word_report(hs_data_list, df_okr, df_rev, period):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for i, hs in enumerate(hs_data_list):
        p = doc.add_heading(f"PHIẾU ĐÁNH GIÁ OKR - {period}", 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Học sinh: {hs['HoTen']} - Lớp: {hs['Lop']} | Email: {hs['Email']}")
        doc.add_paragraph("-" * 60)
        
        doc.add_heading('I. KẾT QUẢ OKR', level=1)
        sub_okr = df_okr[(df_okr['Email'] == hs['Email']) & (df_okr['Dot'] == period)]
        
        if not sub_okr.empty:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            headers = ['Mục Tiêu', 'KR', 'Đích', 'Đạt', '%', 'PH chấm']
            for j, h in enumerate(headers): hdr[j].text = h
            
            for _, row in sub_okr.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(row['MucTieu'])
                cells[1].text = str(row['KetQuaThenChot'])
                cells[2].text = f"{row['MucTieuSo']} {row['DonVi']}"
                cells[3].text = str(row['ThucDat'])
                cells[4].text = f"{row['TienDo']:.1f}%"
                stars = int(row['DiemHaiLong_PH'])
                cells[5].text = "★" * stars if stars > 0 else "-"
        else:
            doc.add_paragraph("(Chưa có dữ liệu OKR)")

        doc.add_heading('II. TỔNG KẾT & PHẢN HỒI', level=1)
        sub_rev = df_rev[(df_rev['Email'] == hs['Email']) & (df_rev['Dot'] == period)]
        
        gv_cmt, ph_cmt, status = "", "", "Chưa chốt"
        if not sub_rev.empty:
            r = sub_rev.iloc[0]
            gv_cmt = r['NhanXet_CuoiKy']
            ph_cmt = r['PhanHoi_PH']
            status = r['TrangThai_CuoiKy']

        doc.add_paragraph(f"1. Nhận xét của GVCN ({status}):")
        doc.add_paragraph(gv_cmt if gv_cmt else "...")
        doc.add_paragraph(f"2. Ý kiến của Gia đình:")
        doc.add_paragraph(ph_cmt if ph_cmt else "...")
        
        if i < len(hs_data_list) - 1:
            doc.add_page_break()
            
    bio = BytesIO()
    doc.save(bio)
    return bio

# =============================================================================
# 4. AUTH & SIDEBAR
# =============================================================================

def login_ui():
    st.markdown("<h1 style='text-align: center;'>🏫 HỆ THỐNG OKR (SUPABASE)</h1>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        with st.form("login_form"):
            email = st.text_input("Email")
            password = st.text_input("Mật khẩu", type="password")
            submit = st.form_submit_button("Đăng nhập", use_container_width=True)
            
            if submit:
                # Master Admin
                if email == "admin@school.com" and password == "123":
                    st.session_state.user = {'Email': email, 'Role': 'Admin', 'HoTen': 'Super Admin'}
                    cookie_manager.set("user_email", email, expires_at=datetime.datetime.now() + datetime.timedelta(days=7))
                    st.rerun()
                
                df = load_data('Users')
                if df.empty:
                    st.error("Chưa có dữ liệu Users.")
                    return
                
                # Check Normal User
                match = df[(df['Email'] == email) & (df['Password'] == password)]
                if not match.empty:
                    st.session_state.user = match.iloc[0].to_dict()
                    cookie_manager.set("user_email", email, expires_at=datetime.datetime.now() + datetime.timedelta(days=7))
                    st.rerun()
                
                # Check Parent
                ph_match = df[(df['EmailPH'] == email) & (df['Password'] == password)]
                if not ph_match.empty:
                    child = ph_match.iloc[0]
                    st.session_state.user = {
                        'Email': email, 'Role': 'PhuHuynh',
                        'HoTen': f"PH em {child['HoTen']}",
                        'ChildEmail': child['Email'], 'ChildName': child['HoTen']
                    }
                    cookie_manager.set("user_email", email, expires_at=datetime.datetime.now() + datetime.timedelta(days=7))
                    st.rerun()
                
                st.error("Sai thông tin đăng nhập.")

def sidebar_controller():
    with st.sidebar:
        st.info(f"👤 {st.session_state.user['HoTen']}\nRole: {st.session_state.user['Role']}")
        
        # Period Selector
        df_p = load_data('Periods')
        if df_p.empty:
            st.warning("Chưa có đợt nào.")
            if st.button("Đăng xuất"):
                cookie_manager.delete("user_email")
                st.session_state.user = None
                st.rerun()
            return None, False
        
        p_opts = df_p['TenDot'].tolist()
        # Default to first open period
        idx = 0
        opens = df_p[df_p['TrangThai'] == 'Mở']['TenDot'].tolist()
        if opens and opens[0] in p_opts:
            idx = p_opts.index(opens[0])
            
        sel_period = st.selectbox("📅 Đợt đánh giá:", p_opts, index=idx)
        status = df_p[df_p['TenDot'] == sel_period].iloc[0]['TrangThai']
        is_open = (status == 'Mở')
        
        if is_open: st.success(f"Trạng thái: {status} 🟢")
        else: st.error(f"Trạng thái: {status} 🔒")
        
        if st.button("🚪 Đăng xuất", use_container_width=True):
            cookie_manager.delete("user_email")
            st.session_state.user = None
            st.rerun()
            
        return sel_period, is_open

# =============================================================================
# 5. MODULES: ADMIN
# =============================================================================

def admin_view(period, is_open):
    st.title("🛡️ Admin Dashboard")
    t1, t2, t3 = st.tabs(["⚙️ Quản lý Đợt", "📊 Thống kê", "👨‍🏫 Giáo Viên"])
    
    with t1:
        with st.form("new_period"):
            c1, c2 = st.columns([3, 1])
            np = c1.text_input("Tên đợt mới")
            if c2.form_submit_button("Tạo"):
                upsert_data('Periods', {'TenDot': np, 'TrangThai': 'Mở'})
                st.success("Tạo thành công!")
                st.rerun()
        
        df_p = load_data('Periods')
        for i, r in df_p.iterrows():
            c1, c2, c3 = st.columns([2, 1, 1])
            c1.write(f"**{r['TenDot']}**")
            c2.write(r['TrangThai'])
            new_stt = "Khóa" if r['TrangThai'] == "Mở" else "Mở"
            if c3.button(f"Đổi sang {new_stt}", key=f"p_{i}"):
                upsert_data('Periods', {'TenDot': r['TenDot'], 'TrangThai': new_stt})
                st.rerun()

    with t2:
        st.markdown(f"**Thống kê Đợt: {period}**")
        df_u = load_data('Users')
        df_okr = load_data('OKRs')
        df_okr = df_okr[df_okr['Dot'] == period]
        
        df_gv = df_u[df_u['Role'] == 'GiaoVien']
        stats = []
        for _, gv in df_gv.iterrows():
            lop = gv['Lop']
            siso = int(gv['SiSo'])
            okrs_cls = df_okr[df_okr['Lop'] == lop]
            submitted = okrs_cls['Email'].nunique()
            approved = okrs_cls[okrs_cls['TrangThai'] == 'Đã duyệt']['Email'].nunique()
            
            stats.append({
                "Lớp": lop, "GVCN": gv['HoTen'], "Sĩ số": siso,
                "Đã nộp": f"{submitted} ({int(submitted/siso*100) if siso else 0}%)",
                "Đã duyệt": f"{approved} ({int(approved/siso*100) if siso else 0}%)"
            })
        st.dataframe(pd.DataFrame(stats), use_container_width=True)

    with t3:
        df_gv = load_data('Users')
        df_gv = df_gv[df_gv['Role'] == 'GiaoVien']
        st.dataframe(df_gv[['Email', 'HoTen', 'Lop', 'SiSo']])
        
        c1, c2 = st.columns(2)
        with c1:
            with st.form("add_gv"):
                e = st.text_input("Email")
                n = st.text_input("Tên")
                l = st.text_input("Lớp")
                s = st.number_input("Sĩ số", 0)
                if st.form_submit_button("Lưu"):
                    upsert_data('Users', {'Email': e, 'Password': '123', 'Role': 'GiaoVien', 'HoTen': n, 'Lop': l, 'SiSo': s})
                    st.success("Lưu thành công!")
                    st.rerun()
        with c2:
            del_email = st.text_input("Nhập Email GV cần xóa:")
            if st.button("Xóa GV"):
                if delete_data('Users', 'Email', del_email):
                    st.success("Đã xóa!")
                    st.rerun()

# =============================================================================
# 6. MODULE: TEACHER (ALL-IN-ONE & ACTIVE EXPANDER FIX)
# =============================================================================

def teacher_view(period, is_open):
    user = st.session_state.user
    my_class = str(user.get('Lop', ''))
    st.title(f"👩‍🏫 LỚP {my_class}")
    
    # Load Data
    df_users = load_data('Users')
    df_hs = df_users[(df_users['Role'] == 'HocSinh') & (df_users['Lop'] == my_class)]
    df_okr = load_data('OKRs')
    df_okr_class = df_okr[(df_okr['Lop'] == my_class) & (df_okr['Dot'] == period)]
    df_rev = load_data('FinalReviews')
    df_rev_class = df_rev[(df_rev['Dot'] == period)]

    t1, t2, t3 = st.tabs(["🚀 Duyệt & Đánh Giá", "👥 Quản Lý HS", "🖨️ Báo Cáo"])

    # --- TAB 1: DUYỆT BÀI (FIX LAG EXPANDER) ---
    with t1:
        if df_hs.empty: st.info("Chưa có học sinh.")
        else:
            for idx, hs in df_hs.iterrows():
                email_hs = hs['Email']
                name_hs = hs['HoTen']
                
                # Check status
                hs_okrs = df_okr_class[df_okr_class['Email'] == email_hs]
                hs_rev = df_rev_class[df_rev_class['Email'] == email_hs]
                
                icon = "🔴"
                status_txt = "Chưa nộp"
                if not hs_okrs.empty:
                    approved = len(hs_okrs[hs_okrs['TrangThai'] == 'Đã duyệt'])
                    if approved == len(hs_okrs): icon, status_txt = "🟢", "Đã duyệt OKR"
                    else: icon, status_txt = "🟡", "Chờ duyệt OKR"
                
                is_finalized = False
                if not hs_rev.empty and hs_rev.iloc[0]['TrangThai_CuoiKy'] == 'Đã chốt':
                    icon, status_txt = "✅", "Đã chốt sổ"
                    is_finalized = True
                elif not hs_rev.empty:
                    icon, status_txt = "⏳", "Đang đánh giá"

                # EXPANDER LOGIC
                is_expanded = (st.session_state.active_expander == email_hs)
                with st.expander(f"{icon} {name_hs} ({status_txt})", expanded=is_expanded):
                    
                    # 1. OKR SECTION
                    st.markdown("**1. OKR**")
                    if hs_okrs.empty: st.warning("Trống")
                    else:
                        for _, row in hs_okrs.iterrows():
                            c1, c2, c3 = st.columns([3, 1.5, 1.5])
                            c1.markdown(f"- {row['MucTieu']} / **{row['KetQuaThenChot']}**")
                            c1.caption(f"Đạt: {row['ThucDat']}/{row['MucTieuSo']} {row['DonVi']}")
                            
                            color = "green" if row['TrangThai'] == 'Đã duyệt' else "orange"
                            c2.markdown(f":{color}[{row['TrangThai']}]")
                            
                            if is_open:
                                if row['YeuCauXoa'] == 'TRUE':
                                    c3.error("Xin xóa!")
                                    if c3.button("Xóa", key=f"del_{row['ID']}"):
                                        st.session_state.active_expander = email_hs # Keep open
                                        delete_data('OKRs', 'ID', row['ID'])
                                        st.rerun()
                                else:
                                    if row['TrangThai'] != 'Đã duyệt' and c3.button("Duyệt", key=f"app_{row['ID']}"):
                                        st.session_state.active_expander = email_hs
                                        upsert_data('OKRs', {'ID': row['ID'], 'TrangThai': 'Đã duyệt'})
                                        st.rerun()
                                    if row['TrangThai'] != 'Cần sửa' and c3.button("Sửa", key=f"rej_{row['ID']}"):
                                        st.session_state.active_expander = email_hs
                                        upsert_data('OKRs', {'ID': row['ID'], 'TrangThai': 'Cần sửa'})
                                        st.rerun()
                        st.divider()

                    # 2. REVIEW SECTION
                    st.markdown("**2. Đánh giá Cuối kỳ**")
                    curr_txt = hs_rev.iloc[0]['NhanXet_CuoiKy'] if not hs_rev.empty else ""
                    ph_fb = hs_rev.iloc[0]['PhanHoi_PH'] if not hs_rev.empty else "Chưa có."
                    st.info(f"PH: {ph_fb}")
                    
                    with st.form(f"rev_{email_hs}"):
                        txt = st.text_area("Nhận xét GV:", value=curr_txt, disabled=not is_open)
                        fin = st.checkbox("Chốt sổ?", value=is_finalized, disabled=not is_open)
                        if st.form_submit_button("Lưu"):
                            st.session_state.active_expander = email_hs
                            stt_val = "Đã chốt" if fin else "Chưa chốt"
                            upsert_data('FinalReviews', {'Email': email_hs, 'Dot': period, 'NhanXet_CuoiKy': txt, 'TrangThai_CuoiKy': stt_val})
                            st.success("Lưu thành công")
                            time.sleep(0.5)
                            st.rerun()

    # --- TAB 2: QUẢN LÝ HS ---
    with t2:
        c1, c2 = st.columns(2)
        with c1:
            st.dataframe(df_hs[['Email', 'HoTen', 'EmailPH']])
            del_hs = st.text_input("Email HS cần xóa:")
            if st.button("Xóa HS"):
                if delete_data('Users', 'Email', del_hs):
                    st.success("Đã xóa")
                    st.rerun()
        with c2:
            with st.form("add_hs"):
                e = st.text_input("Email")
                n = st.text_input("Tên")
                p = st.text_input("Email PH")
                if st.form_submit_button("Thêm"):
                    upsert_data('Users', {'Email': e, 'Password': '123', 'Role': 'HocSinh', 'HoTen': n, 'Lop': my_class, 'EmailPH': p, 'SiSo': 0})
                    st.success("Đã thêm")
                    st.rerun()

    # --- TAB 3: BÁO CÁO ---
    with t3:
        if st.button("Tải Báo Cáo Cả Lớp (.docx)"):
            hs_data = df_hs.to_dict('records')
            bio = generate_word_report(hs_data, df_okr, df_rev, period)
            st.download_button("Download", bio, f"OKR_{my_class}.docx")

# =============================================================================
# 7. MODULE: STUDENT
# =============================================================================

def student_view(period, is_open):
    user = st.session_state.user
    st.title(f"🎓 {user['HoTen']}")
    
    df_okr = load_data('OKRs')
    my_okrs = df_okr[(df_okr['Email'] == user['Email']) & (df_okr['Dot'] == period)]
    df_rev = load_data('FinalReviews')
    rev = df_rev[(df_rev['Email'] == user['Email']) & (df_rev['Dot'] == period)]

    # 1. REVIEW
    st.markdown("### 📝 Tổng kết")
    gv_txt = rev.iloc[0]['NhanXet_CuoiKy'] if not rev.empty and rev.iloc[0]['NhanXet_CuoiKy'] else "Chưa có."
    ph_txt = rev.iloc[0]['PhanHoi_PH'] if not rev.empty and rev.iloc[0]['PhanHoi_PH'] else "Chưa có."
    st.info(f"**GV Nhận xét:** {gv_txt}")
    st.warning(f"**PH Phản hồi:** {ph_txt}")
    st.divider()

    # 2. CREATE
    if is_open:
        with st.expander("➕ Thêm OKR Mới"):
            with st.form("new_okr"):
                o = st.text_input("Mục tiêu")
                k = st.text_input("Key Result")
                t = st.number_input("Mục tiêu số", min_value=0.0)
                u = st.text_input("Đơn vị")
                if st.form_submit_button("Lưu"):
                    uid = str(uuid.uuid4())
                    upsert_data('OKRs', {
                        'ID': uid, 'Email': user['Email'], 'Lop': user['Lop'], 'Dot': period,
                        'MucTieu': o, 'KetQuaThenChot': k, 'MucTieuSo': t, 'DonVi': u,
                        'ThucDat': 0.0, 'TienDo': 0.0, 'TrangThai': 'Chờ duyệt', 'YeuCauXoa': 'FALSE'
                    })
                    st.success("Đã thêm")
                    st.rerun()

    # 3. LIST & UPDATE
    st.subheader("Tiến độ")
    if my_okrs.empty: st.info("Chưa có OKR")
    else:
        objs = my_okrs['MucTieu'].unique()
        for obj in objs:
            with st.container(border=True):
                st.markdown(f"**🎯 {obj}**")
                krs = my_okrs[my_okrs['MucTieu'] == obj]
                for _, row in krs.iterrows():
                    st.divider()
                    st.markdown(f"{row['KetQuaThenChot']} ({row['TrangThai']})")
                    c1, c2, c3 = st.columns([2, 3, 1])
                    c1.caption(f"Đích: {row['MucTieuSo']} {row['DonVi']}")
                    
                    cur_act = float(row['ThucDat'])
                    if is_open and row['TrangThai'] == 'Đã duyệt':
                        new_act = c2.number_input(f"Đạt ({row['DonVi']})", value=cur_act, step=0.01, key=f"act_{row['ID']}")
                        prog = calculate_progress(new_act, row['MucTieuSo'])
                        c2.progress(int(prog))
                        if c3.button("Lưu", key=f"up_{row['ID']}"):
                            upsert_data('OKRs', {'ID': row['ID'], 'ThucDat': new_act, 'TienDo': prog})
                            st.success("Updated")
                            st.rerun()
                    else:
                        c2.write(f"Đạt: {cur_act}")
                        c2.progress(int(row['TienDo']))
                    
                    if is_open and row['YeuCauXoa'] == 'FALSE':
                        if c3.button("Xin xóa", key=f"dx_{row['ID']}"):
                            upsert_data('OKRs', {'ID': row['ID'], 'YeuCauXoa': 'TRUE'})
                            st.rerun()

# =============================================================================
# 8. MODULE: PARENT
# =============================================================================

def parent_view(period, is_open):
    user = st.session_state.user
    st.title(f"👨‍👩‍👧‍👦 PHHS: {user['ChildName']}")
    
    df_okr = load_data('OKRs')
    child_okrs = df_okr[(df_okr['Email'] == user['ChildEmail']) & (df_okr['Dot'] == period)]
    
    st.subheader("Đánh giá OKR")
    if child_okrs.empty: st.info("Con chưa có OKR")
    else:
        for _, row in child_okrs.iterrows():
            with st.container(border=True):
                c1, c2 = st.columns([3, 1])
                c1.write(f"**KR:** {row['KetQuaThenChot']}")
                c1.caption(f"Tiến độ: {row['TienDo']}%")
                cur_star = int(row['DiemHaiLong_PH'])
                new_star = c2.slider(f"Sao ({row['ID']})", 1, 5, cur_star if cur_star > 0 else 5)
                if c2.button("Lưu sao", key=f"s_{row['ID']}"):
                    upsert_data('OKRs', {'ID': row['ID'], 'DiemHaiLong_PH': new_star})
                    st.success("Đã lưu")

    st.divider()
    df_rev = load_data('FinalReviews')
    rev = df_rev[(df_rev['Email'] == user['ChildEmail']) & (df_rev['Dot'] == period)]
    gv_txt = rev.iloc[0]['NhanXet_CuoiKy'] if not rev.empty else "Chưa có"
    st.info(f"GV Nhận xét: {gv_txt}")
    
    ph_old = rev.iloc[0]['PhanHoi_PH'] if not rev.empty else ""
    with st.form("ph_f"):
        txt = st.text_area("Ý kiến gia đình:", value=ph_old)
        if st.form_submit_button("Gửi"):
            upsert_data('FinalReviews', {'Email': user['ChildEmail'], 'Dot': period, 'PhanHoi_PH': txt})
            st.success("Đã gửi")
            st.rerun()

# =============================================================================
# 9. MAIN RUN
# =============================================================================

def main():
    # Auto Login Check
    if not st.session_state.user:
        cookie_email = cookie_manager.get(cookie="user_email")
        if cookie_email:
            df = load_data('Users')
            if not df.empty:
                # Normal User
                match = df[df['Email'] == cookie_email]
                if not match.empty:
                    st.session_state.user = match.iloc[0].to_dict()
                else:
                    # Parent
                    ph_match = df[df['EmailPH'] == cookie_email]
                    if not ph_match.empty:
                        child = ph_match.iloc[0]
                        st.session_state.user = {
                            'Email': cookie_email, 'Role': 'PhuHuynh',
                            'HoTen': f"PH em {child['HoTen']}",
                            'ChildEmail': child['Email'], 'ChildName': child['HoTen']
                        }
            # Admin Master Bypass
            if cookie_email == "admin@school.com":
                st.session_state.user = {'Email': cookie_email, 'Role': 'Admin', 'HoTen': 'Super Admin'}

    if not st.session_state.user:
        login_ui()
    else:
        period, is_open = sidebar_controller()
        if not period:
            st.warning("Vui lòng liên hệ Admin tạo đợt.")
            return
        
        role = st.session_state.user['Role']
        if role == 'Admin': admin_view(period, is_open)
        elif role == 'GiaoVien': teacher_view(period, is_open)
        elif role == 'HocSinh': student_view(period, is_open)
        elif role == 'PhuHuynh': parent_view(period, is_open)

if __name__ == "__main__":
    main()
